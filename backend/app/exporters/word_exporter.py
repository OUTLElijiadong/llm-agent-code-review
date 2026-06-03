"""
Word报告导出器: 使用python-docx生成审查报告
"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def export_word_report(detail: dict) -> BytesIO:
    """将报告详情数据渲染为Word文档

    Args:
        detail: 报告详情字典(由report_service.get_report_detail返回)

    Returns:
        BytesIO: Word文档的二进制缓冲
    """
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(11)

    title = doc.add_heading("代码审查报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    project = detail.get("project", {})
    task = detail.get("task", {})
    doc.add_paragraph(f"项目: {project.get('project_name', project.get('name', ''))}")
    doc.add_paragraph(f"语言: {project.get('language', '')}")
    doc.add_paragraph(f"审查时间: {task.get('create_time', '')}")
    doc.add_paragraph(f"综合评分: {task.get('score', 0)}/100")

    stats = detail.get("stats", {})
    doc.add_heading("审查统计", level=1)
    doc.add_paragraph(f"审查文件数: {stats.get('total_files', 0)}")
    doc.add_paragraph(f"发现问题总数: {stats.get('total_issues', 0)}")

    severity = stats.get("severity", {})
    doc.add_paragraph(f"严重: {severity.get('严重', 0)}  高: {severity.get('高', 0)}  "
                      f"中: {severity.get('中', 0)}  低: {severity.get('低', 0)}")

    by_type = stats.get("by_type", {})
    if by_type:
        doc.add_heading("问题类型分布", level=1)
        for t, c in by_type.items():
            doc.add_paragraph(f"{t}: {c}个")

    summary = detail.get("summary", "")
    if summary:
        doc.add_heading("总体评价", level=1)
        doc.add_paragraph(summary)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

"""Word 报告导出服务 (T12)

基于 python-docx 构建 .docx 文档,支持中文字体与严重度颜色编码
(严重:红 / 高:橙 / 中:黄 / 低:蓝),修复代码使用 Courier New 字体。

依赖:
- T11 report_exporter.export_to_dict: 构建报告上下文(task_info/issues/statistics)
- python-docx >= 1.1.0: Word 文档生成

文档结构:
    1. 报告头(标题 / 任务名 / 审查时间 / 综合评分)
    2. 统计摘要表(问题总数 / 严重度分布 / 已修复数)
    3. 问题列表(每个问题一个段落,按严重度分组)
"""
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.report_exporter import export_to_dict

# ============ 模块常量 ============

# 严重度 → RGB 颜色映射(python-docx RGBColor)
_SEVERITY_RGB: Dict[str, RGBColor] = {
    "严重": RGBColor(0xDC, 0x26, 0x26),  # 红色
    "高": RGBColor(0xEA, 0x58, 0x0C),    # 橙色
    "中": RGBColor(0xCA, 0x8A, 0x04),    # 黄色
    "低": RGBColor(0x25, 0x63, 0xEB),    # 蓝色
}

# 严重度展示顺序(严重 > 高 > 中 > 低 > 其他)
_SEVERITY_ORDER: Tuple[str, ...] = ("严重", "高", "中", "低")

# 中文字体名称(SimSun 是 Word 常用中文字体)
_CHINESE_FONT: str = "SimSun"

# 代码字体名称(Courier New 等宽字体,适合展示代码)
_CODE_FONT: str = "Courier New"


# ============ 内部辅助 ============

def _severity_rgb(severity: Optional[str]) -> RGBColor:
    """获取严重度对应的 RGB 颜色。

    Args:
        severity: 严重度字符串(严重/高/中/低)。

    Returns:
        RGBColor: 对应的 RGB 颜色对象;未知严重度返回灰色。
    """
    if severity and severity in _SEVERITY_RGB:
        return _SEVERITY_RGB[severity]
    return RGBColor(0x6B, 0x72, 0x80)  # 灰色


def _set_run_chinese_font(run, font_name: str) -> None:
    """设置 run 的中文字体(包括东亚字体)。

    python-docx 默认只设置西文字体,中文需要额外设置 w:eastAsia 属性。

    Args:
        run: docx.run.Run 对象。
        font_name: 字体名称(如 "SimSun")。

    Returns:
        None
    """
    run.font.name = font_name
    # 设置东亚字体(中文渲染必需)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _add_run(paragraph, text: str, *, bold: bool = False, color: Optional[RGBColor] = None,
             font_name: str = _CHINESE_FONT, size: int = 11) -> Any:
    """向段落添加一个 run 并设置格式。

    Args:
        paragraph: docx 段落对象。
        text: run 文本内容。
        bold: 是否加粗,默认 False。
        color: 文字颜色(RGBColor),默认 None(黑色)。
        font_name: 字体名称,默认 SimSun(中文)。
        size: 字号(pt),默认 11。

    Returns:
        Any: 创建的 run 对象。
    """
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    _set_run_chinese_font(run, font_name)
    return run


def _group_issues_by_severity(issues: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    """按严重度分组 issue(严重 > 高 > 中 > 低 > 其他)。

    Args:
        issues: 已归一化的 issue 字典列表。

    Returns:
        Dict[str, List[Dict[str, Any]]]: 键为严重度,值为该严重度下的 issue 列表。
    """
    grouped: Dict[str, List[Dict[str, Any]]] = {key: [] for key in _SEVERITY_ORDER}
    grouped["其他"] = []
    for issue in issues or []:
        severity = issue.get("severity")
        if severity in grouped:
            grouped[severity].append(issue)
        else:
            grouped["其他"].append(issue)
    return grouped


def _format_score(score: int) -> str:
    """格式化评分展示文本(含风险等级)。

    Args:
        score: 综合评分 0-100。

    Returns:
        str: 格式化后的评分文本,如 "72/100 (中风险)"。
    """
    if score >= 80:
        level = "低风险"
    elif score >= 60:
        level = "中风险"
    elif score >= 40:
        level = "高风险"
    else:
        level = "极高风险"
    return f"{score}/100 ({level})"


# ============ Word 文档构建 ============

def _build_header(doc: Any, task_info: Dict[str, Any], score: int) -> None:
    """构建报告头部(标题 + 任务元信息)。

    Args:
        doc: python-docx Document 对象。
        task_info: 任务信息字典。
        score: 综合评分。

    Returns:
        None
    """
    # 主标题(居中)
    title_para = doc.add_heading(level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title_para, "代码审查报告", bold=True, size=22)

    # 任务元信息
    meta_fields = [
        ("任务名称", str(task_info.get("task_name") or "未命名任务")),
        ("所属项目", str(task_info.get("project_name") or "-")),
        ("审查类型", str(task_info.get("review_type") or "-")),
        ("审查模型", str(task_info.get("model_name") or "-")),
        ("生成时间", str(task_info.get("create_time") or "-")),
    ]
    for label, value in meta_fields:
        para = doc.add_paragraph()
        _add_run(para, f"{label}: ", bold=True)
        _add_run(para, value)

    # 综合评分(加粗显示)
    score_para = doc.add_paragraph()
    _add_run(score_para, "综合评分: ", bold=True)
    _add_run(score_para, _format_score(score), bold=True, color=RGBColor(0x11, 0x18, 0x27))


def _build_summary_table(doc: Any, statistics: Dict[str, Any]) -> None:
    """构建统计摘要表格。

    Args:
        doc: python-docx Document 对象。
        statistics: 统计信息字典。

    Returns:
        None
    """
    doc.add_heading("统计摘要", level=1)

    severity_count = statistics.get("severity_count", {})
    rows_data = [
        ("指标", "数值"),
        ("问题总数", str(statistics.get("total_issues", 0))),
        ("严重问题", str(severity_count.get("严重", 0))),
        ("高优先级问题", str(severity_count.get("高", 0))),
        ("中优先级问题", str(severity_count.get("中", 0))),
        ("低优先级问题", str(severity_count.get("低", 0))),
        ("已修复问题", str(statistics.get("fixed_count", 0))),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, (label, value) in enumerate(rows_data):
        cells = table.rows[row_idx].cells
        # 标签列
        label_para = cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(label_para, label, bold=(row_idx == 0), size=10)
        # 数值列
        value_para = cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(value_para, value, bold=(row_idx == 0), size=10)

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Pt(120)
        row.cells[1].width = Pt(80)


def _build_issue_paragraph(doc: Any, issue: Dict[str, Any], index: int, severity: str) -> None:
    """构建单个 issue 的段落(标题 + 基本信息 + 描述 + 修复建议 + 修复代码)。

    Args:
        doc: python-docx Document 对象。
        issue: 已归一化的 issue 字典。
        index: issue 在当前严重度分组中的序号(从 1 开始)。
        severity: 严重度字符串。

    Returns:
        None
    """
    color = _severity_rgb(severity)

    # 问题标题(严重度色块 + 序号 + 标题)
    title_para = doc.add_paragraph()
    _add_run(title_para, f"[{severity}] ", bold=True, color=color, size=12)
    title_text = str(issue.get("title") or "未命名问题")
    _add_run(title_para, f"{index}. {title_text}", bold=True, size=12)

    # 基本信息
    info_para = doc.add_paragraph()
    file_name = str(issue.get("file_name") or "-")
    line_number = issue.get("line_number") or "-"
    issue_type = str(issue.get("issue_type") or "-")
    cwe = str(issue.get("cwe") or "-")
    _add_run(info_para, f"文件:{file_name} | 行号:{line_number} | 类型:{issue_type} | CWE:{cwe}", size=10)

    # 问题描述
    description = issue.get("description") or ""
    if description:
        desc_para = doc.add_paragraph()
        _add_run(desc_para, "描述:", bold=True, size=10)
        _add_run(desc_para, str(description), size=10)

    # 修复建议
    suggestion = issue.get("suggestion") or ""
    if suggestion:
        sug_para = doc.add_paragraph()
        _add_run(sug_para, "修复建议:", bold=True, size=10)
        _add_run(sug_para, str(suggestion), size=10)

    # 修复代码(Courier New 字体)
    fixed_code = issue.get("fixed_code") or ""
    if fixed_code:
        code_label_para = doc.add_paragraph()
        _add_run(code_label_para, "修复代码:", bold=True, size=10)
        # 代码内容使用 Courier New 等宽字体
        code_para = doc.add_paragraph()
        _add_run(code_para, str(fixed_code), font_name=_CODE_FONT, size=9,
                 color=RGBColor(0x05, 0x96, 0x69))


def _build_issues_section(doc: Any, issues: List[Dict[str, Any]]) -> None:
    """构建问题列表章节(按严重度分组)。

    Args:
        doc: python-docx Document 对象。
        issues: 已归一化且按严重度排序的 issue 字典列表。

    Returns:
        None
    """
    doc.add_heading("问题详情", level=1)

    grouped = _group_issues_by_severity(issues)
    has_any_issue = any(len(grouped[key]) > 0 for key in grouped)

    if not has_any_issue:
        empty_para = doc.add_paragraph()
        _add_run(empty_para, "本次审查未发现问题。", size=11)
        return

    for severity in _SEVERITY_ORDER + ("其他",):
        group = grouped.get(severity, [])
        if not group:
            continue
        color = _severity_rgb(severity)
        # 分组标题
        group_para = doc.add_paragraph()
        _add_run(group_para, f"{severity}级问题({len(group)} 个)", bold=True, color=color, size=12)

        for idx, issue in enumerate(group, start=1):
            _build_issue_paragraph(doc, issue, idx, severity)


# ============ 对外导出接口 ============

def export_to_word(
    task: Any,
    issues: List[Any],
    summary: Optional[str],
    score: int,
    template_type: str = "detailed",
) -> bytes:
    """导出 Word 报告字节流。

    使用 python-docx 构建 .docx 文档,包含报告头、统计摘要表、
    按严重度分组的问题列表。中文字体使用 SimSun,修复代码使用 Courier New。

    Args:
        task: 审查任务(TaskDetailOut / ReviewTask / dict)。
        issues: issue 对象列表(IssueOut / ReviewIssue / dict)。
        summary: AI 总体评价文本,可为 None。
        score: 综合评分 0-100。
        template_type: 模板类型(simple/detailed/compliance),目前仅影响
            章节详略预留,Word 内容结构一致。默认 "detailed"。

    Returns:
        bytes: Word 文档的二进制字节流(可直接写入文件或作为 HTTP 响应体)。
    """
    # 构建报告上下文(复用 T11 的归一化逻辑)
    context = export_to_dict(task, issues, summary, score)
    task_info = context.get("task_info", {})
    statistics = context.get("statistics", {})
    sorted_issues = context.get("issues", [])
    summary_text = context.get("summary") or ""

    # 创建 Word 文档
    doc = Document()

    # 设置默认样式字体
    normal_style = doc.styles["Normal"]
    normal_style.font.name = _CHINESE_FONT
    normal_style.font.size = Pt(11)
    # 设置默认样式的东亚字体
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), _CHINESE_FONT)

    # 1. 报告头
    _build_header(doc, task_info, score)

    # 2. 总体评价
    if summary_text:
        doc.add_heading("总体评价", level=1)
        summary_para = doc.add_paragraph()
        _add_run(summary_para, str(summary_text), size=11)

    # 3. 统计摘要表
    _build_summary_table(doc, statistics)

    # 4. 问题详情(按严重度分组)
    _build_issues_section(doc, sorted_issues)

    # 导出字节流
    buffer = BytesIO()
    doc.save(buffer)
    word_bytes = buffer.getvalue()
    buffer.close()
    return word_bytes

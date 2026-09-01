"""代码分片、编码转换与 Word 导出覆盖率补全测试。

这些用例只验证纯逻辑和内存文件，不依赖数据库、网络或真实磁盘输出。
"""
from __future__ import annotations

import base64
from io import BytesIO

import chardet
import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.ai import code_chunker
from app.exporters.word_exporter import export_word_report
from app.utils import encoding_utils


def _document_text(document: Document) -> str:
    """合并 Word 文档的全部段落文本。

    Args:
        document: 已解析的 python-docx 文档。

    Returns:
        str: 以换行拼接的段落正文。
    """
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_chunk_code_returns_single_chunk_below_threshold():
    """短代码应保持完整并返回正确的零基行号范围。

    Returns:
        None: 断言空文本和普通短文本的单分片行为。
    """
    empty = code_chunker.chunk_code("", "python")
    assert empty == [code_chunker.CodeChunk(text="", start_line=0, end_line=0)]

    content = "line one\nline two\n"
    chunks = code_chunker.chunk_code(content, "python", threshold=100)
    assert chunks == [code_chunker.CodeChunk(text=content, start_line=0, end_line=2)]


def test_chunk_code_uses_function_boundaries_and_preserves_file_header():
    """超长 Python 代码应按定义边界切分并保留首个定义前的头部。

    Returns:
        None: 断言边界发现、相邻定义合并和阈值截断。
    """
    content = (
        "# module header\n"
        "import os\n"
        "\n"
        "def first():\n"
        "    return 1\n"
        "\n"
        "def second():\n"
        "    return 2\n"
        "\n"
        "class LargeBlock:\n"
        f"    value = '{'x' * 80}'\n"
    )
    boundaries = code_chunker._find_function_boundaries(
        content,
        code_chunker._FUNC_PATTERNS["python"],
    )
    assert boundaries == [3, 6, 9]

    chunks = code_chunker.chunk_code(content, "python", threshold=60)
    assert chunks[0].text == "# module header\nimport os\n\n"
    assert (chunks[0].start_line, chunks[0].end_line) == (0, 3)
    assert "def first" in chunks[1].text
    assert "def second" in chunks[1].text
    assert "class LargeBlock" not in chunks[1].text
    assert "class LargeBlock" in chunks[2].text
    assert "".join(chunk.text for chunk in chunks) == content


def test_chunk_code_falls_back_to_fixed_line_windows():
    """未知语言仍须无损切分，并严格遵守上下文阈值。

    Returns:
        None: 断言 450 行兜底分片和空行列表辅助函数。
    """
    content = "".join(f"line {index}\n" for index in range(450))
    chunks = code_chunker.chunk_code(content, "plaintext", threshold=10)
    assert "".join(chunk.text for chunk in chunks) == content
    assert all(len(chunk.text) <= 10 for chunk in chunks)
    assert chunks[0].start_line == 0
    assert chunks[-1].end_line == 450

    no_boundary_content = "value = 1\n" * 250
    no_boundaries = code_chunker.chunk_code(no_boundary_content, "python", threshold=10)
    assert "".join(chunk.text for chunk in no_boundaries) == no_boundary_content
    assert all(len(chunk.text) <= 10 for chunk in no_boundaries)
    assert code_chunker._split_by_lines([], lines_per_chunk=200) == []


def test_to_utf8_handles_empty_utf8_and_binary_bytes():
    """编码转换应直返 UTF-8 文本，并将含空字节内容标记为 Base64。

    Returns:
        None: 断言空输入、中文 UTF-8 和二进制往返解码。
    """
    assert encoding_utils.to_utf8(b"") == ""
    assert encoding_utils.to_utf8("你好，世界".encode("utf-8")) == "你好，世界"
    assert encoding_utils._is_binary(b"") is False
    assert encoding_utils._is_binary(b"plain text") is False
    assert encoding_utils._is_binary(b"png\x00payload") is True

    raw = b"\x89PNG\x00\x01\x02"
    encoded = encoding_utils.to_utf8(raw)
    assert encoded.startswith(encoding_utils.BASE64_PREFIX)
    payload = encoded[len(encoding_utils.BASE64_PREFIX):]
    assert base64.b64decode(payload) == raw
    assert encoding_utils._encode_base64(raw) == encoded


def test_to_utf8_uses_detected_encoding_or_low_confidence_fallback(monkeypatch):
    """非 UTF-8 文本应采用可信探测结果，低置信度时退回 Base64。

    Args:
        monkeypatch: Pytest 属性替换工具。

    Returns:
        None: 断言 Latin-1 解码和低置信度二进制保护。
    """
    raw = b"caf\xe9"

    def detect_latin_one(_raw: bytes) -> dict:
        """返回高置信度 Latin-1 探测结果。

        Args:
            _raw: 待探测原始字节。

        Returns:
            dict: chardet 兼容的编码与置信度结果。
        """
        return {"encoding": "iso-8859-1", "confidence": 0.99}

    monkeypatch.setattr(chardet, "detect", detect_latin_one)
    assert encoding_utils.to_utf8(raw) == "café"

    def detect_low_confidence(_raw: bytes) -> dict:
        """返回低置信度探测结果。

        Args:
            _raw: 待探测原始字节。

        Returns:
            dict: 低于业务阈值的探测结果。
        """
        return {"encoding": None, "confidence": 0.1}

    monkeypatch.setattr(chardet, "detect", detect_low_confidence)
    fallback = encoding_utils.to_utf8(raw)
    assert fallback == encoding_utils.BASE64_PREFIX + base64.b64encode(raw).decode("ascii")


def test_to_utf8_falls_back_when_detected_codec_is_invalid(monkeypatch):
    """探测到不可用编码器时应捕获异常并安全保存 Base64。

    Args:
        monkeypatch: Pytest 属性替换工具。

    Returns:
        None: 断言解码器查找异常不会向上传播。
    """
    raw = b"\xff\xfeinvalid"

    def detect_invalid_codec(_raw: bytes) -> dict:
        """返回不存在的高置信度编码名称。

        Args:
            _raw: 待探测原始字节。

        Returns:
            dict: 会令 bytes.decode 抛出 LookupError 的结果。
        """
        return {"encoding": "codec-does-not-exist", "confidence": 1.0}

    monkeypatch.setattr(chardet, "detect", detect_invalid_codec)
    assert encoding_utils.to_utf8(raw).startswith(encoding_utils.BASE64_PREFIX)


def test_export_word_report_empty_detail_is_valid_and_readable():
    """空报告数据也应生成可重新打开且包含默认统计的有效 DOCX。

    Returns:
        None: 断言缓冲区位置、ZIP 文件头、默认文本和基础样式。
    """
    buffer = export_word_report({})
    assert isinstance(buffer, BytesIO)
    assert buffer.tell() == 0
    assert buffer.getvalue()[:2] == b"PK"

    document = Document(buffer)
    text = _document_text(document)
    assert "代码审查报告" in text
    assert "项目: " in text
    assert "语言: " in text
    assert "审查时间: " in text
    assert "综合评分: 0/100" in text
    assert "审查文件数: 0" in text
    assert "发现问题总数: 0" in text
    assert "严重: 0  高: 0  中: 0  低: 0" in text
    assert document.styles["Normal"].font.name == "SimSun"
    assert document.styles["Normal"].font.size.pt == pytest.approx(11)
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_export_word_report_renders_project_stats_types_and_summary():
    """完整报告应渲染项目名称、分级统计、类型分布和总体评价。

    Returns:
        None: 通过重新读取 DOCX 正文验证内容而非仅检查非空字节。
    """
    detail = {
        "project": {"project_name": "安全审查平台", "name": "备用名称", "language": "Python"},
        "task": {"create_time": "2026-07-10 12:30:00", "score": 92},
        "stats": {
            "total_files": 12,
            "total_issues": 6,
            "severity": {"严重": 1, "高": 2, "中": 2, "低": 1},
            "by_type": {"安全漏洞": 3, "性能问题": 2, "代码规范": 1},
        },
        "summary": "整体质量良好，建议优先修复高风险问题。",
    }
    buffer = export_word_report(detail)
    document = Document(buffer)
    text = _document_text(document)

    assert "项目: 安全审查平台" in text
    assert "备用名称" not in text
    assert "语言: Python" in text
    assert "审查时间: 2026-07-10 12:30:00" in text
    assert "综合评分: 92/100" in text
    assert "审查文件数: 12" in text
    assert "发现问题总数: 6" in text
    assert "严重: 1  高: 2  中: 2  低: 1" in text
    assert "问题类型分布" in text
    assert "安全漏洞: 3个" in text
    assert "性能问题: 2个" in text
    assert "代码规范: 1个" in text
    assert "总体评价" in text
    assert "整体质量良好，建议优先修复高风险问题。" in text


def test_export_word_report_uses_legacy_project_name_fallback():
    """缺少 project_name 时应兼容旧版 project.name 字段。

    Returns:
        None: 断言旧字段回退且空类型/摘要不会生成对应章节。
    """
    buffer = export_word_report({
        "project": {"name": "旧版项目", "language": "Java"},
        "task": {"score": 70},
        "stats": {"by_type": {}, "severity": {}},
        "summary": "",
    })
    text = _document_text(Document(buffer))
    assert "项目: 旧版项目" in text
    assert "语言: Java" in text
    assert "问题类型分布" not in text
    assert "总体评价" not in text
